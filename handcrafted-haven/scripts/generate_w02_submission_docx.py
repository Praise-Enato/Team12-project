from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "doc" / "W02-Group-Project-Submission.docx"


def set_default_font(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    paragraph.space_after = Pt(6)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.15


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    set_default_font(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("W02 Group Project")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.paragraph_format.space_after = Pt(14)

    add_heading(document, "1. Summary")
    add_body(
        document,
        "Our group met to review the Handcrafted Haven project requirements "
        "and begin planning the structure of the web application. During the "
        "meeting, we discussed the purpose of the marketplace, the core features "
        "required for sellers and shoppers, the overall design direction, and "
        "the first set of user stories that can guide development. We agreed that "
        "the site should present a warm, trustworthy, artisan-focused experience "
        "while still feeling modern and easy to use. We also identified the need "
        "for responsive design, accessibility, strong navigation, and a clear "
        "product browsing experience as major priorities for the project."
    )
    add_body(
        document,
        "Participants in the meeting included [Your Name], [Teammate 1], "
        "[Teammate 2], and [Teammate 3].",
    )

    add_heading(document, "2. URL")
    add_body(
        document,
        "The URL of the group project's GitHub repository is: "
        "[Paste your GitHub repository link here]",
    )

    add_heading(document, "3. Local Repo Project")
    add_body(
        document,
        "I created a local copy of the repository on my computer and used it to "
        "begin organizing the planning documents for the project. My local "
        "project folder is located at "
        "/Users/mac/Team12-project/handcrafted-haven. Evidence of the local "
        "repository setup includes the project README, the design planning "
        "document, and the project board seed document that were created inside "
        "the local project folder.",
    )

    add_heading(document, "4. Design and Styling Identify")
    add_body(
        document,
        "We identified the initial design and styling direction for Handcrafted "
        "Haven as part of our planning process. The project will use a warm, "
        "handcrafted visual theme with earthy colors and a clean layout that "
        "supports product discovery and seller storytelling. Our planned color "
        "scheme includes Clay (#A35A3A), Forest (#355C4D), Cream (#F6F0E8), "
        "Walnut (#4A3428), and Sand (#DCC7AA). For typography, we selected "
        "Cormorant Garamond for headings to give the site a refined, artisan-"
        "market feel and Manrope for body text and interface elements to keep "
        "the site readable and modern. We also identified other styling "
        "elements such as rounded buttons, soft card shadows, product-focused "
        "layouts, responsive navigation, and accessibility-conscious contrast "
        "and focus states.",
    )

    add_heading(document, "5. User Stories and Work Items")
    add_body(
        document,
        "Our group created an initial list of more than ten work items to begin "
        "organizing the project board and development effort. These work items "
        "include seller account signup and login, seller profile creation and "
        "editing, product listing management, product image upload, marketplace "
        "browsing, filtering and sorting, product detail pages, ratings and "
        "written reviews, shopping cart and checkout flow, responsive navigation "
        "and mobile experience, accessibility and usability review, and "
        "deployment with team workflow setup. These items provide a starting "
        "point for the GitHub Project board and will continue to evolve as the "
        "group refines the project over the next two weeks.",
    )

    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
